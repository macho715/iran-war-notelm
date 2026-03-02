"""Build .docx summary from Upgrade / Deep / Verifier reports. Requires: pip install python-docx"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for j, h in enumerate(headers):
        t.rows[0].cells[j].text = h
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            if j < len(t.rows[i + 1].cells):
                t.rows[i + 1].cells[j].text = str(cell)
    return t

def main():
    out_dir = Path(__file__).resolve().parent.parent / "docs"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "Upgrade_Deep_Verifier_Summary.docx"

    doc = Document()
    doc.add_heading("Upgrade / Deep / Verifier 요약 보고서", 0)
    doc.add_paragraph("Iran-UAE Monitor project-upgrade 산출물 요약 (2026-03-02)")

    # --- 1. Upgrade Report ---
    add_heading(doc, "1. Upgrade Report 요약", 1)
    add_para(doc, "Executive Summary", bold=True)
    add_para(doc, "현재: Iran-UAE 실시간 모니터링 - Python(Playwright/httpx/NotebookLM) 스크래퍼 + Phase 2 AI + Telegram/WhatsApp, Option A JSON 아카이브, SQLite/Postgres(Phase 4). Docker/GHA hourly 지원.")
    add_para(doc, "문서 감사: README/ARCHITECTURE/LAYOUT/CHANGELOG 기준. CONTRIBUTING/SECURITY/ADR 없음. v1.6 통합 운영.")
    add_para(doc, "권장: 보수(A) 30일 Observability/테스트/문서 -> 중간(B) 60일 스케줄러/채널 -> 공격(C) 90일 Phase 4 SSOT/배포.")

    add_para(doc, "Current State Snapshot", bold=True)
    add_table(doc,
        ["Area", "Status", "Risk"],
        [
            ["Backend", "Python 3.11, Playwright, httpx, NotebookLM MCP", "Medium"],
            ["Scheduler", "APScheduler 매 1시간", "Medium"],
            ["Reporting", "Telegram + Twilio WhatsApp", "Low"],
            ["Storage", "SQLite 기본, Postgres(Phase 4)", "Medium"],
            ["CI/CD", "GHA hourly, Python 3.11, Playwright", "Medium"],
            ["Observability", ".health_state.json, health.py", "Medium"],
            ["Deployment", "로컬/Docker, GHA run_now", "High"],
        ],
    )

    add_para(doc, "Upgrade Ideas Top 10 (요약)", bold=True)
    add_table(doc,
        ["Rank", "Idea", "Bucket", "PriorityScore", "First PR"],
        [
            ["1", "Observability: 헬스/메트릭 계약 + 실패 알림", "Reliability/Observability", "4.0", "health + alert"],
            ["2", "Playwright async-only + asyncio.gather 일원화", "Performance", "3.0", "scrapers"],
            ["3", "APScheduler 이벤트 모니터링", "Reliability/Observability", "4.0", "main.py listeners"],
            ["4", "CONTRIBUTING.md + SECURITY.md 초안", "Docs/Process", "15.0", "docs"],
            ["5", "Single-instance lock + stale recovery 테스트", "Reliability/Observability", "15.0", "tests"],
            ["6", "Playwright timeout/networkidle 명시", "DX/Tooling", "12.0", "scrapers"],
            ["7", "Phase 4: Postgres SSOT + GHA 정리", "Architecture", "1.25", "schema + workflow"],
            ["8", "RSS 타임아웃/재시도/User-Agent 정리", "Reliability", "12.0", "config + rss_feed"],
            ["9", "Phase 3: Slack/Email 채널 스캐폴드", "DX/Tooling", "2.0", "reporter extensions"],
            ["10", "Docker 이미지 경량화", "Security/Performance", "3.0", "Dockerfile"],
        ],
    )

    add_para(doc, "Options A/B/C", bold=True)
    add_para(doc, "A(보수) 30일: Observability, Playwright timeout/async, CONTRIBUTING/SECURITY, lock 테스트. 리스크 낮음.")
    add_para(doc, "B(중간) 60일: A + APScheduler 모니터링, RSS 정리, Slack/Email 스캐폴드. 리스크 중간.")
    add_para(doc, "C(공격) 90일: B + Phase 4 Postgres SSOT, GHA 정리, Docker 경량화. 리스크/비용 최대.")

    add_para(doc, "30/60/90 Roadmap", bold=True)
    add_table(doc, ["30d", "60d", "90d"], [
        ["health 스키마 + 실패 알림", "APScheduler 리스너 + 알림", "Postgres SSOT + GHA 정리"],
        ["Playwright async + timeout", "CONTRIBUTING/SECURITY", "Slack/Email 스캐폴드"],
        ["lock guard 테스트", "RSS 타임아웃/재시도", "Docker 멀티스테이지"],
        ["CONTRIBUTING/SECURITY 초안", "Phase 3 채널 설계", "24h 운영 검증"],
    ])

    # --- 2. Deep Report ---
    add_heading(doc, "2. Deep Report 요약 (upgrade-deep-synth)", 1)
    add_para(doc, "Best3 Gate Summary", bold=True)
    add_table(doc,
        ["Best#", "Idea", "PriorityScore", "EvidenceCount", "DateOK", "Final"],
        [
            ["1", "APScheduler 이벤트 모니터링", "4.0", "2", "Yes", "PASS"],
            ["2", "Playwright async + timeout/networkidle", "3.0~12.0", "2", "AMBER", "AMBER"],
            ["3", "Observability: 헬스/메트릭 + 실패 알림", "4.0", "2", "AMBER", "AMBER"],
        ],
    )
    add_para(doc, "Best3 Deep Dive 요약", bold=True)
    add_para(doc, "BEST #1 (APScheduler): 리스너 등록 + 알림 플래그, PR1~3, Mock 테스트, 쿨다운으로 알림 과다 방지.")
    add_para(doc, "BEST #2 (Playwright): async_playwright + goto timeout/networkidle, uae_media -> social_media -> 상수 추출, domcontentloaded 폴백.")
    add_para(doc, "BEST #3 (Observability): health 스키마 + 실패 알림, HEALTH_ALERT_ENABLED 플래그, .health_state.json 하위 호환.")

    # --- 3. Verifier Report ---
    add_heading(doc, "3. Verifier Report 요약 (Deep2 Gate Review)", 1)
    add_para(doc, "PASS/FAIL (Best3)", bold=True)
    add_table(doc,
        ["Idea", "Tier", "Verdict", "Why"],
        [
            ["APScheduler 이벤트 모니터링", "Best3", "PASS", "Evidence>=2, 날짜 충족, PR/Tests/Rollout/KPIs 정의"],
            ["Playwright async + timeout/networkidle", "Best3", "AMBER", "1건 published_date 없음, Deep Dive 완전"],
            ["Observability 헬스/메트릭 + 실패 알림", "Best3", "AMBER", "repo+공식 docs 날짜 없음, 스키마 하위 호환 확인"],
        ],
    )
    add_para(doc, "Apply Gates (0~4)", bold=True)
    add_para(doc, "Gate 0: Dry-run (설정/스키마만). Gate 1: Change list (main.py, config, health, reporter, scrapers, tests). Gate 2: Explicit approval (APPROVE_*). Gate 3: Canary (SCHEDULER_ALERT_ENABLED, HEALTH_ALERT_ENABLED). Gate 4: Rollback (플래그 false, 리스너 제거, 스키마 하위 호환).")
    add_para(doc, "Rollout triggers: Gate 0~2 통과, 테스트 그린, 승인. 스테이징/GHA 1회 확인 후 프로덕션.")
    add_para(doc, "Rollback triggers: 알림 스팸, 스크래퍼 성공률 하락(<90%), 헬스 오류율 상승, GHA 실패가 Best3 변경 때문일 때.")
    add_para(doc, "Final Go/No-Go", bold=True)
    add_para(doc, "결론: GO. APScheduler PASS, Playwright/Observability AMBER(날짜만 미충족). Apply Gates 0~4, Rollback, Test Matrix, PR Sanity 충족. 단계적 플래그/테스트로 적용 권장.")

    doc.add_paragraph()
    add_para(doc, "참조: docs/UPGRADE_REPORT_project-upgrade.md, docs/DEEP_REPORT_upgrade-deep-synth.md, docs/VERIFIER_REPORT_deep2.md")

    doc.save(out_path)
    print(f"Saved: {out_path}")
    return str(out_path)

if __name__ == "__main__":
    main()
